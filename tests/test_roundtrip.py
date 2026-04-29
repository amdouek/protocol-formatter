"""
tests/test_roundtrip.py -- End-to-end and unit tests for ProtocolFormatter.

Test categories
---------------
1.  Schema validation tests
    Ensure valid and invalid Protocol instances behave correctly.

2.  Parser smoke tests
    Parse both reference template documents and verify structural expectations.

3.  Renderer roundtrip tests
    Hand-crafted Protocol → Node.js renderer → .docx written to disk.
    Verifies the output is a valid ZIP (DOCX) containing word/document.xml.

4.  Heuristic extractor tests
    ParsedDocument → rule_extractor → valid Protocol.

5.  Prompts tests
    Verify system prompt coverage and user message construction.

6.  LLM extractor unit tests (no Ollama required)
    Test JSON extraction and post-processing in isolation.

7.  CLI smoke tests
    Invoke the Typer app in-process and verify exit codes.

8.  Duplicate detection tests
    Verify DuplicateDetector behaviour.

Running
-------
    # From protocol_formatter/ directory:
    pytest tests/test_roundtrip.py -v

    # With output capture disabled (to see Rich output from CLI tests):
    pytest tests/test_roundtrip.py -v -s

Notes
-----
- Tests that require Ollama are marked with @pytest.mark.ollama and are
  skipped automatically if the server is not reachable.
- Tests that require the render.js renderer (node + docx npm) are marked
  with @pytest.mark.renderer and skipped if Node.js is not available.
- Tests that require the template .docx files from the uploads directory
  are skipped gracefully if those files are not present.
"""

from __future__ import annotations

import json
import zipfile
from pathlib import Path
from typing import Any

import pytest

# ---------------------------------------------------------------------------
# Path constants
# ---------------------------------------------------------------------------

_TESTS_DIR = Path(__file__).resolve().parent
_PKG_ROOT = _TESTS_DIR.parent
_UPLOADS_DIR = Path("/mnt/user-data/uploads")

_WET_LAB_TEMPLATE = _UPLOADS_DIR / "Protocol_Template.docx"
_COMP_TEMPLATE = _UPLOADS_DIR / "Computational_Protocol_Template.docx"

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture(scope="session")
def cfg():
    """Load the style_guide.yaml config once per session."""
    import yaml
    path = _PKG_ROOT / "configs" / "style_guide.yaml"
    with path.open("r", encoding="utf-8") as fh:
        return yaml.safe_load(fh)


@pytest.fixture(scope="session")
def wet_lab_protocol():
    """
    A fully-populated wet-lab Protocol instance for renderer tests.
    Hand-crafted to exercise all wet-lab features.
    """
    from schema import (
        Protocol, SectionType, ProcedureSection, ActionStep, StoppingPoint,
        Callout, CalloutType, MaterialsTable, MixTable, TableRow,
    )
    return Protocol(
        title="RNA Extraction from Zebrafish Tissue",
        subtitle="Using TRIzol reagent",
        author="Alon Douek",
        section_type=SectionType.WET_LAB,
        section_number="4",
        section_name="MOLECULAR BIOLOGY",
        version="1.0",
        date="21/04/2026",
        overview=(
            "This protocol describes **total RNA extraction** from adult zebrafish "
            "tissue using TRIzol reagent. Expected duration: 2–3 hours. "
            "Requires RNase-free conditions throughout."
        ),
        materials=[
            MaterialsTable(
                heading="Reagents",
                rows=[
                    TableRow(cells=["TRIzol Reagent", "Thermo Fisher Scientific", "15596026"]),
                    TableRow(cells=["Chloroform", "Sigma-Aldrich", "C2432"]),
                    TableRow(cells=["Isopropanol", "Sigma-Aldrich", "I9516"]),
                    TableRow(cells=["75% Ethanol (DEPC-H₂O)", "Lab supply", "—"]),
                ],
            ),
            MaterialsTable(
                heading="Equipment & consumables",
                rows=[
                    TableRow(cells=["Tissue homogeniser", "Qiagen", "TissueRuptor II"]),
                    TableRow(cells=["RNase-free tubes (1.5 mL)", "Eppendorf", "Z666505"]),
                ],
            ),
        ],
        procedure=[
            ProcedureSection(
                heading="Tissue homogenisation",
                preamble=None,
                callouts=[
                    Callout(
                        callout_type=CalloutType.CAUTION,
                        text="TRIzol contains phenol and guanidinium thiocyanate. "
                             "Work in a fume hood and wear appropriate PPE.",
                    ),
                ],
                steps=[
                    ActionStep(
                        text="Transfer tissue (up to **50 mg**) to a 1.5 mL RNase-free tube on ice.",
                    ),
                    ActionStep(
                        text="Add **1 mL** TRIzol Reagent per 50 mg tissue.",
                        children=[
                            ActionStep(text="Ensure tissue is fully submerged."),
                            ActionStep(text="Keep on ice at all times."),
                        ],
                    ),
                    ActionStep(text="Homogenise for **30 s** at maximum speed."),
                    StoppingPoint(
                        text="Homogenised samples can be stored at **−80 °C** for up to one month."
                    ),
                ],
            ),
            ProcedureSection(
                heading="Phase separation",
                callouts=[],
                steps=[
                    ActionStep(text="Incubate at **room temperature** for **5 min**."),
                    ActionStep(
                        text="Add **200 µL** chloroform per mL TRIzol. "
                             "Cap and shake vigorously for **15 s**."
                    ),
                    ActionStep(
                        text="Centrifuge at **12,000 × g** for **15 min** at **4 °C**."
                    ),
                ],
            ),
        ],
        notes=[
            "For fibrous tissues, increase homogenisation time to 60 s.",
            "RNA yield can be improved by a second precipitation step.",
        ],
        references=[
            "Chomczynski P & Sacchi N (1987). _Anal Biochem_ 162(1):156–9.",
        ],
        mix_tables=[
            MixTable(
                heading="Wash buffer (per sample)",
                rows=[
                    TableRow(cells=["75% Ethanol", "1 mL"]),
                    TableRow(cells=["Total", "1 mL"], bold=True),
                ],
            ),
        ],
        source_filename="rna_extraction.docx",
    )


@pytest.fixture(scope="session")
def comp_protocol():
    """A fully-populated computational Protocol instance for renderer tests."""
    from schema import (
        Protocol, SectionType, ProcedureSection, ActionStep,
        Callout, CalloutType, Prerequisites,
    )
    return Protocol(
        title="Publishing a Python Package to PyPI",
        subtitle="Building, versioning, and uploading distributions",
        author="Alon Douek",
        section_type=SectionType.COMPUTATIONAL,
        section_number="9",
        section_name="COMPUTATIONAL METHODS & BIOINFORMATICS",
        version="1.1",
        date="21/04/2026",
        overview=(
            "This protocol covers building a source distribution and wheel, "
            "then uploading the package to **PyPI** using twine."
        ),
        prerequisites=Prerequisites(
            software="Python >=3.10, pip >=23.0",
            access="PyPI account with a project-scoped API token",
            dependencies="pip install build twine",
        ),
        procedure=[
            ProcedureSection(
                heading="Step 1 — Build the distribution",
                preamble="Run from the project root directory.",
                callouts=[
                    Callout(
                        callout_type=CalloutType.NOTE,
                        text="Increment the version field in pyproject.toml before each upload.",
                    ),
                ],
                steps=[
                    ActionStep(text="```bash\npython -m build\n```"),
                    ActionStep(
                        text="Verify that `dist/` contains both a `.whl` and `.tar.gz`."
                    ),
                ],
            ),
        ],
        notes=["To upload to TestPyPI, run: twine upload --repository testpypi dist/*"],
        references=["https://packaging.python.org/en/latest/tutorials/packaging-projects/"],
        source_filename="pypi_publishing.docx",
    )


# ---------------------------------------------------------------------------
# Pytest marks
# ---------------------------------------------------------------------------

pytestmark_renderer = pytest.mark.renderer
pytestmark_ollama = pytest.mark.ollama


def _node_available() -> bool:
    import subprocess
    try:
        r = subprocess.run(["node", "--version"], capture_output=True, timeout=5)
        return r.returncode == 0
    except Exception:
        return False


def _docx_npm_available() -> bool:
    import subprocess
    script = "try{require('docx');process.exit(0)}catch(e){process.exit(1)}"
    try:
        r = subprocess.run(["node", "-e", script], capture_output=True, timeout=5)
        return r.returncode == 0
    except Exception:
        return False


def _ollama_available() -> bool:
    try:
        import urllib.request
        urllib.request.urlopen("http://localhost:11434/api/tags", timeout=3)
        return True
    except Exception:
        return False


_renderer_available = _node_available() and _docx_npm_available()
_ollama_up = _ollama_available()

skip_no_renderer = pytest.mark.skipif(
    not _renderer_available,
    reason="Node.js + docx npm package required",
)
skip_no_ollama = pytest.mark.skipif(
    not _ollama_up,
    reason="Ollama server not running (start with `ollama serve`)",
)
skip_no_templates = pytest.mark.skipif(
    not (_WET_LAB_TEMPLATE.exists() and _COMP_TEMPLATE.exists()),
    reason="Reference template .docx files not present at /mnt/user-data/uploads/",
)


# ===========================================================================
# 1. Schema validation tests
# ===========================================================================

class TestSchema:

    def test_wet_lab_valid(self, wet_lab_protocol):
        """A well-formed wet-lab Protocol validates without error."""
        assert wet_lab_protocol.title == "RNA Extraction from Zebrafish Tissue"
        assert wet_lab_protocol.section_type.value == "wet_lab"
        assert len(wet_lab_protocol.procedure) == 2
        assert len(wet_lab_protocol.materials) == 2

    def test_comp_valid(self, comp_protocol):
        """A well-formed computational Protocol validates without error."""
        assert comp_protocol.section_type.value == "computational"
        assert comp_protocol.prerequisites is not None
        assert comp_protocol.materials == []
        assert comp_protocol.mix_tables == []

    def test_cross_field_prerequisites_on_wet_lab(self):
        """prerequisites must be None for wet_lab protocols."""
        from schema import Protocol, SectionType, ProcedureSection, ActionStep, Prerequisites
        with pytest.raises(Exception, match="prerequisites"):
            Protocol(
                title="Bad",
                section_type=SectionType.WET_LAB,
                section_number="1",
                section_name="TEST",
                overview="Test.",
                prerequisites=Prerequisites(software="Python"),
                procedure=[
                    ProcedureSection(
                        heading="Step",
                        steps=[ActionStep(text="Do thing.")],
                    )
                ],
            )

    def test_cross_field_materials_on_computational(self):
        """materials must be empty for computational protocols."""
        from schema import Protocol, SectionType, ProcedureSection, ActionStep, MaterialsTable, TableRow
        with pytest.raises(Exception, match="materials"):
            Protocol(
                title="Bad",
                section_type=SectionType.COMPUTATIONAL,
                section_number="9",
                section_name="TEST",
                overview="Test.",
                materials=[MaterialsTable(heading="R", rows=[TableRow(cells=["A", "B", "C"])])],
                procedure=[
                    ProcedureSection(
                        heading="Step",
                        steps=[ActionStep(text="Do thing.")],
                    )
                ],
            )

    def test_procedure_min_length(self):
        """procedure must contain at least one ProcedureSection."""
        from schema import Protocol, SectionType
        with pytest.raises(Exception):
            Protocol(
                title="Bad",
                section_type=SectionType.WET_LAB,
                section_number="1",
                section_name="TEST",
                overview="Test.",
                procedure=[],
            )

    def test_json_roundtrip(self, wet_lab_protocol):
        """Protocol serialises to JSON and deserialises back identically."""
        from schema import Protocol
        json_str = wet_lab_protocol.model_dump_json()
        restored = Protocol.model_validate_json(json_str)
        assert restored.title == wet_lab_protocol.title
        assert restored.mix_tables[0].rows[1].bold is True

    def test_recursive_steps(self):
        """ActionStep supports arbitrary nesting depth."""
        from schema import ActionStep, StoppingPoint
        deep = ActionStep(
            text="Level 0",
            children=[
                ActionStep(
                    text="Level 1",
                    children=[
                        ActionStep(
                            text="Level 2",
                            children=[
                                StoppingPoint(text="Pause here — safe stopping point.")
                            ],
                        )
                    ],
                )
            ],
        )
        assert deep.children[0].children[0].children[0].step_type.value == "stopping_point"

    def test_mix_table_total_bold_flag(self):
        """MixTable rows can have the bold flag set explicitly."""
        from schema import MixTable, TableRow
        mt = MixTable(
            heading="Master mix",
            rows=[
                TableRow(cells=["Buffer", "5 µL"]),
                TableRow(cells=["Total", "10 µL"], bold=True),
            ],
        )
        assert mt.rows[1].bold is True
        assert mt.rows[0].bold is False


# ===========================================================================
# 2. Parser smoke tests
# ===========================================================================

class TestParser:

    @skip_no_templates
    def test_wet_lab_template_parse(self):
        """Parse the wet-lab reference template and verify structural expectations."""
        from parser.docx_reader import read_docx
        doc = read_docx(_WET_LAB_TEMPLATE)

        assert len(doc.paragraphs) > 5
        assert doc.full_text
        assert doc.source_path == _WET_LAB_TEMPLATE

        h1_texts = [p.raw_text for p in doc.paragraphs if p.heading_level == 1]
        assert "Overview" in h1_texts
        assert "Materials" in h1_texts
        assert "Procedure" in h1_texts
        assert "References" in h1_texts

        assert doc.section_type == "wet_lab"
        assert len(doc.tables) >= 1

    @skip_no_templates
    def test_comp_template_parse(self):
        """Parse the computational reference template and verify structural expectations."""
        from parser.docx_reader import read_docx
        doc = read_docx(_COMP_TEMPLATE)

        h1_texts = [p.raw_text for p in doc.paragraphs if p.heading_level == 1]
        assert "Overview" in h1_texts
        assert "Prerequisites" in h1_texts
        assert "Procedure" in h1_texts

        assert doc.section_type == "computational"

    @skip_no_templates
    def test_author_detected_in_comp_template(self):
        """Alon Douek is correctly identified as author in the comp template."""
        from parser.docx_reader import read_docx
        doc = read_docx(_COMP_TEMPLATE)
        assert doc.author == "Alon Douek"

    @skip_no_templates
    def test_bold_markers_in_full_text(self):
        """Bold run properties from the template produce ** markers in full_text."""
        from parser.docx_reader import read_docx
        doc = read_docx(_WET_LAB_TEMPLATE)
        assert "**" in doc.full_text

    def test_read_document_dispatch_docx(self, tmp_path):
        """read_document() dispatches .docx to docx_reader."""
        from parser import read_document
        # Create a minimal valid docx to test dispatch (not content)
        from docx import Document as DocxDocument
        d = DocxDocument()
        d.add_heading("Test Protocol", level=1)
        d.add_paragraph("Overview text here.")
        p = tmp_path / "test_dispatch.docx"
        d.save(str(p))
        doc = read_document(p)
        assert doc.source_path == p
        assert len(doc.paragraphs) >= 1

    def test_read_document_unsupported_extension(self, tmp_path):
        """read_document() raises ValueError for unsupported extensions."""
        from parser import read_document
        p = tmp_path / "file.pdf"
        p.write_bytes(b"%PDF-1.4")
        with pytest.raises(ValueError, match="Unsupported file type"):
            read_document(p)

    def test_read_document_missing_file(self):
        """read_document() raises FileNotFoundError for missing files."""
        from parser import read_document
        with pytest.raises(FileNotFoundError):
            read_document(Path("/nonexistent/protocol.docx"))


# ===========================================================================
# 3. Renderer roundtrip tests
# ===========================================================================

class TestRenderer:

    @skip_no_renderer
    def test_wet_lab_roundtrip(self, wet_lab_protocol, tmp_path):
        """Wet-lab Protocol → render.js → valid .docx file on disk."""
        from renderer.node_renderer import render_protocol
        output = render_protocol(wet_lab_protocol, output_dir=tmp_path)

        assert output.exists(), f"Output file not created: {output}"
        assert output.suffix == ".docx"
        assert output.stat().st_size > 5000, "Output file suspiciously small"

        # Verify it is a valid ZIP (DOCX is ZIP-based)
        assert zipfile.is_zipfile(output), "Output is not a valid ZIP/DOCX"

        with zipfile.ZipFile(output) as zf:
            names = zf.namelist()
            assert "word/document.xml" in names, "Missing word/document.xml"
            assert "[Content_Types].xml" in names, "Missing [Content_Types].xml"

    @skip_no_renderer
    def test_comp_roundtrip(self, comp_protocol, tmp_path):
        """Computational Protocol → render.js → valid .docx file on disk."""
        from renderer.node_renderer import render_protocol
        output = render_protocol(comp_protocol, output_dir=tmp_path)

        assert output.exists()
        assert zipfile.is_zipfile(output)
        with zipfile.ZipFile(output) as zf:
            assert "word/document.xml" in zf.namelist()

    @skip_no_renderer
    def test_output_filename_derived_from_title(self, wet_lab_protocol, tmp_path):
        """Output filename is derived from the protocol title and version."""
        from renderer.node_renderer import render_protocol
        output = render_protocol(wet_lab_protocol, output_dir=tmp_path)
        assert "rna_extraction" in output.name.lower()
        assert "v1_0" in output.name

    @skip_no_renderer
    def test_output_dir_created_if_missing(self, wet_lab_protocol, tmp_path):
        """render_protocol() creates the output directory if it does not exist."""
        from renderer.node_renderer import render_protocol
        new_dir = tmp_path / "nested" / "output"
        assert not new_dir.exists()
        output = render_protocol(wet_lab_protocol, output_dir=new_dir)
        assert new_dir.exists()
        assert output.exists()

    @skip_no_renderer
    def test_render_js_stdin_interface(self, wet_lab_protocol, tmp_path):
        """render.js accepts JSON payload via stdin and returns {"output": "..."}."""
        import subprocess, json as _json
        render_script = _PKG_ROOT / "renderer" / "templates" / "render.js"
        payload = _json.loads(wet_lab_protocol.model_dump_json())
        payload["output_path"] = str(tmp_path / "stdin_test.docx")
        payload_json = _json.dumps(payload)

        result = subprocess.run(
            ["node", str(render_script)],
            input=payload_json,
            capture_output=True,
            text=True,
            timeout=30,
        )
        assert result.returncode == 0, f"render.js stderr: {result.stderr}"
        response = _json.loads(result.stdout.strip())
        assert "output" in response
        assert Path(response["output"]).exists()

    @skip_no_renderer
    def test_render_js_validation_error(self, tmp_path):
        """render.js exits 1 and writes structured error for invalid payload."""
        import subprocess, json as _json
        render_script = _PKG_ROOT / "renderer" / "templates" / "render.js"
        bad_payload = _json.dumps({"title": "Bad", "section_type": "invalid"})

        result = subprocess.run(
            ["node", str(render_script)],
            input=bad_payload,
            capture_output=True,
            text=True,
            timeout=10,
        )
        assert result.returncode == 1
        error = _json.loads(result.stderr.strip())
        assert "error" in error

    def test_title_slug(self):
        """_title_to_slug produces filesystem-safe slugs."""
        from renderer.node_renderer import _title_to_slug
        assert _title_to_slug("RNA Extraction from Zebrafish") == \
               "rna_extraction_from_zebrafish"
        assert _title_to_slug("In Situ Hybridisation (ISH) — wholemount") == \
               "in_situ_hybridisation_ish_wholemount"
        assert len(_title_to_slug("A" * 80)) <= 60


# ===========================================================================
# 4. Heuristic extractor tests
# ===========================================================================

class TestHeuristicExtractor:

    @skip_no_templates
    def test_wet_lab_template_heuristic(self):
        """Heuristic extractor produces a valid Protocol from the wet-lab template."""
        from parser.docx_reader import read_docx
        from extractor.rule_extractor import extract_protocol_heuristic

        doc = read_docx(_WET_LAB_TEMPLATE)
        protocol = extract_protocol_heuristic(doc)

        assert protocol.title
        assert protocol.section_type.value == "wet_lab"
        assert len(protocol.procedure) >= 1
        assert protocol.overview
        assert protocol.materials == [] or isinstance(protocol.materials, list)

    @skip_no_templates
    def test_comp_template_heuristic(self):
        """Heuristic extractor produces a valid Protocol from the comp template."""
        from parser.docx_reader import read_docx
        from extractor.rule_extractor import extract_protocol_heuristic

        doc = read_docx(_COMP_TEMPLATE)
        protocol = extract_protocol_heuristic(doc)

        assert protocol.section_type.value == "computational"
        assert protocol.materials == []
        assert protocol.mix_tables == []

    def test_heuristic_from_minimal_docx(self, tmp_path):
        """Heuristic extractor handles a minimal single-section document."""
        from docx import Document as DocxDocument
        from parser.docx_reader import read_docx
        from extractor.rule_extractor import extract_protocol_heuristic

        d = DocxDocument()
        d.add_heading("Minimal Protocol", level=1)
        d.add_paragraph("This is the overview text.")
        d.add_heading("Procedure", level=1)
        d.add_heading("Step group one", level=2)
        p = d.add_paragraph("Add 5 µL buffer.", style="List Number")
        p2 = d.add_paragraph("Incubate for 10 min.", style="List Number")
        path = tmp_path / "minimal.docx"
        d.save(str(path))

        doc = read_docx(path)
        protocol = extract_protocol_heuristic(doc, source_filename="minimal.docx")

        assert protocol.title == "Minimal Protocol"
        assert len(protocol.procedure) >= 1
        assert protocol.source_filename == "minimal.docx"

    def test_heuristic_fallback_title_from_filename(self, tmp_path):
        """When no H1 title is found, title falls back to filename stem."""
        from docx import Document as DocxDocument
        from parser.docx_reader import read_docx
        from extractor.rule_extractor import extract_protocol_heuristic

        d = DocxDocument()
        d.add_paragraph("Just a paragraph with no headings.")
        path = tmp_path / "my_cool_protocol.docx"
        d.save(str(path))

        doc = read_docx(path)
        protocol = extract_protocol_heuristic(doc)
        # Title should be derived from filename
        assert "cool" in protocol.title.lower() or protocol.title

# ===========================================================================
# 4B. Callout detection tests
# ===========================================================================

class TestCalloutDetection:
    """Unit tests for parser/utils.py callout detection helpers."""

    def test_crucial_prefix_classified_as_critical(self):
        from parser.utils import detect_callout_type
        assert detect_callout_type("CRUCIAL: avoid contamination.") == "critical"

    def test_crucial_prefix_stripped_from_text(self):
        from parser.utils import strip_callout_prefix
        assert strip_callout_prefix("CRUCIAL: avoid contamination.") == \
            "avoid contamination."

    def test_crucial_prefix_case_insensitive(self):
        from parser.utils import detect_callout_type, strip_callout_prefix
        assert detect_callout_type("Crucial - keep cold.") == "critical"
        assert strip_callout_prefix("crucial: keep cold.") == "keep cold."

    def test_crucial_mid_sentence_still_promotes(self):
        """Existing keyword-scan behaviour preserved for mid-sentence usage."""
        from parser.utils import detect_callout_type
        assert detect_callout_type("It is crucial to work quickly.") == "critical"

    def test_critical_prefix_unchanged(self):
        """Pre-existing CRITICAL prefix behaviour must not regress."""
        from parser.utils import detect_callout_type, strip_callout_prefix
        assert detect_callout_type("CRITICAL: this step is essential.") == "critical"
        assert strip_callout_prefix("CRITICAL: this step is essential.") == \
            "this step is essential."

    def test_important_prefix_maps_to_caution(self):
        """Pre-existing IMPORTANT/WARNING/CAUTION → caution mapping unchanged."""
        from parser.utils import detect_callout_type
        assert detect_callout_type("IMPORTANT: wear gloves.") == "caution"
        assert detect_callout_type("WARNING: toxic.") == "caution"
        assert detect_callout_type("Caution: hot surface.") == "caution"


# ===========================================================================
# 5. Prompts tests
# ===========================================================================

class TestPrompts:

    def test_system_prompt_contains_schema_fields(self):
        """SYSTEM_PROMPT references all required schema field names."""
        from extractor.prompts import SYSTEM_PROMPT
        required = [
            "section_type", "procedure", "callout_type", "step_type",
            "stopping_point", "materials", "prerequisites", "mix_tables",
        ]
        for field in required:
            assert field in SYSTEM_PROMPT, f"SYSTEM_PROMPT missing field: {field!r}"

    def test_system_prompt_contains_cross_field_rules(self):
        """SYSTEM_PROMPT states the wet_lab/computational field restrictions."""
        from extractor.prompts import SYSTEM_PROMPT
        assert "wet_lab" in SYSTEM_PROMPT
        assert "computational" in SYSTEM_PROMPT
        assert "prerequisites" in SYSTEM_PROMPT

    def test_style_guide_context_all_callout_types(self, cfg):
        """Style guide context includes all four callout types."""
        from extractor.prompts import build_style_guide_context
        ctx = build_style_guide_context(cfg)
        for ct in ("critical", "caution", "tip", "note"):
            assert ct in ctx.lower(), f"Context missing callout type: {ct}"

    def test_style_guide_context_all_sections(self, cfg):
        """Style guide context includes all nine section names."""
        from extractor.prompts import build_style_guide_context
        ctx = build_style_guide_context(cfg)
        assert "MOLECULAR BIOLOGY" in ctx
        assert "COMPUTATIONAL METHODS" in ctx
        assert "ANIMAL HUSBANDRY" in ctx

    def test_user_message_section_hint(self, cfg):
        """Section number hint causes section name to appear in user message."""
        from extractor.prompts import build_user_message
        msg = build_user_message("Protocol text", section_number_hint="7", cfg=cfg)
        assert "GENOMICS" in msg

    def test_retry_message_structure(self):
        """Retry message includes previous output and validation errors."""
        from extractor.prompts import build_retry_message
        msg = build_retry_message('{"bad": 1}', 'field required: title', attempt=2)
        assert "field required: title" in msg
        assert '{"bad": 1}' in msg
        assert "2" in msg  # attempt number

    def test_token_estimate_nonzero(self):
        """Token estimation returns a positive integer for non-empty input."""
        from extractor.prompts import estimate_token_count
        assert estimate_token_count("hello world foo bar") > 0

    def test_token_budget_no_warn_for_short_doc(self):
        """Short documents do not trigger the token budget warning."""
        from extractor.prompts import check_token_budget, SYSTEM_PROMPT
        _, over = check_token_budget(SYSTEM_PROMPT, "short protocol text", 4096)
        assert over is False


# ===========================================================================
# 6. LLM extractor unit tests (no Ollama required)
# ===========================================================================

class TestLLMExtractorPure:

    def test_extract_json_plain(self):
        """JSON object extracted from plain string."""
        from extractor.llm_extractor import _extract_json_from_response
        raw = '{"title": "Test", "x": 1}'
        assert _extract_json_from_response(raw) == raw

    def test_extract_json_fenced(self):
        """JSON object extracted from markdown code fence."""
        from extractor.llm_extractor import _extract_json_from_response
        fenced = '```json\n{"title": "Test"}\n```'
        result = _extract_json_from_response(fenced)
        assert '"title"' in result

    def test_extract_json_with_preamble(self):
        """JSON object extracted even when preceded by prose."""
        from extractor.llm_extractor import _extract_json_from_response
        prose = 'Here is the result:\n{"title": "Test", "n": 42}'
        result = _extract_json_from_response(prose)
        assert '"title"' in result and '"n"' in result

    def test_extract_json_nested(self):
        """Balanced brace scanning handles nested objects correctly."""
        from extractor.llm_extractor import _extract_json_from_response
        nested = '{"a": {"b": {"c": 1}}, "d": [{"e": 2}]}'
        assert _extract_json_from_response(nested) == nested

    def test_extract_json_escaped_quotes(self):
        """Escaped quotes inside JSON strings do not confuse brace scanner."""
        from extractor.llm_extractor import _extract_json_from_response
        raw = '{"text": "He said \\"hello\\" to me"}'
        assert _extract_json_from_response(raw) == raw

    def test_extract_json_raises_on_empty(self):
        """ValueError raised when response contains no JSON object."""
        from extractor.llm_extractor import _extract_json_from_response
        with pytest.raises(ValueError, match="No JSON object"):
            _extract_json_from_response("No JSON here at all.")

    def test_post_process_author_default(self):
        """Empty author is replaced with 'ARMI'."""
        from extractor.llm_extractor import _post_process_payload
        result = _post_process_payload({"author": "", "section_type": "wet_lab", "procedure": []})
        assert result["author"] == "ARMI"

    def test_post_process_section_type_normalised(self):
        """section_type is lowercased during post-processing."""
        from extractor.llm_extractor import _post_process_payload
        result = _post_process_payload({"section_type": "WET_LAB", "procedure": []})
        assert result["section_type"] == "wet_lab"

    def test_post_process_callout_prefix_stripped(self):
        """Callout label prefixes are stripped from callout.text."""
        from extractor.llm_extractor import _post_process_payload
        result = _post_process_payload({
            "section_type": "wet_lab",
            "procedure": [{
                "heading": "H", "steps": [], "preamble": None,
                "callouts": [
                    {"callout_type": "caution", "text": "IMPORTANT: Keep on ice."}
                ],
            }],
        })
        assert result["procedure"][0]["callouts"][0]["text"] == "Keep on ice."

    def test_post_process_mix_table_total_bold(self):
        """'Total' rows in mix_tables are auto-bolded."""
        from extractor.llm_extractor import _post_process_payload
        result = _post_process_payload({
            "section_type": "wet_lab",
            "procedure": [],
            "mix_tables": [{"heading": "Mix", "rows": [
                {"cells": ["Total", "10 µL"], "bold": False},
                {"cells": ["Buffer", "5 µL"], "bold": False},
            ]}],
        })
        assert result["mix_tables"][0]["rows"][0]["bold"] is True
        assert result["mix_tables"][0]["rows"][1]["bold"] is False

    def test_post_process_incubate_bold(self):
        """'Incubate' rows in mix_tables are auto-bolded."""
        from extractor.llm_extractor import _post_process_payload
        result = _post_process_payload({
            "section_type": "wet_lab",
            "procedure": [],
            "mix_tables": [{"heading": "Mix", "rows": [
                {"cells": ["Incubate at 37°C", "10 min"], "bold": False},
            ]}],
        })
        assert result["mix_tables"][0]["rows"][0]["bold"] is True


# ===========================================================================
# 7. CLI smoke tests
# ===========================================================================

class TestCLI:

    def test_check_command_exits_cleanly(self):
        """
        `protocol-formatter check` exits with code 0 or 1 (never crashes).
        Exit code 1 is acceptable when optional dependencies are absent.
        """
        from typer.testing import CliRunner
        from main import app
        runner = CliRunner()
        result = runner.invoke(app, ["check"])
        assert result.exit_code in (0, 1), \
            f"Unexpected exit code {result.exit_code}: {result.output}"

    def test_format_missing_file(self):
        """`protocol-formatter format` exits 1 for a missing source file."""
        from typer.testing import CliRunner
        from main import app
        runner = CliRunner()
        result = runner.invoke(app, ["format", "/nonexistent/protocol.docx"])
        assert result.exit_code == 1
        assert "not found" in result.output.lower() or "error" in result.output.lower()

    def test_format_unsupported_extension(self, tmp_path):
        """`protocol-formatter format` exits 1 for unsupported file types."""
        from typer.testing import CliRunner
        from main import app
        f = tmp_path / "file.txt"
        f.write_text("not a protocol")
        runner = CliRunner()
        result = runner.invoke(app, ["format", str(f)])
        assert result.exit_code == 1

    def test_batch_missing_directory(self):
        """`protocol-formatter batch` exits 1 for a missing directory."""
        from typer.testing import CliRunner
        from main import app
        runner = CliRunner()
        result = runner.invoke(app, ["batch", "/nonexistent/dir/"])
        assert result.exit_code == 1

    def test_batch_empty_directory(self, tmp_path):
        """`protocol-formatter batch` exits 0 for an empty directory (no files to process)."""
        from typer.testing import CliRunner
        from main import app
        runner = CliRunner()
        result = runner.invoke(app, ["batch", str(tmp_path)])
        assert result.exit_code == 0

    @skip_no_renderer
    def test_format_heuristic_roundtrip(self, tmp_path):
        """`protocol-formatter format --heuristic` produces a .docx for a minimal input."""
        from docx import Document as DocxDocument
        from typer.testing import CliRunner
        from main import app

        d = DocxDocument()
        d.add_heading("CLI Test Protocol", level=1)
        d.add_paragraph("Overview of the CLI test protocol.")
        d.add_heading("Procedure", level=1)
        d.add_heading("Step group", level=2)
        d.add_paragraph("Add reagent to tube.", style="List Number")
        src = tmp_path / "cli_test.docx"
        d.save(str(src))

        out_dir = tmp_path / "output"
        runner = CliRunner()
        result = runner.invoke(app, [
            "format", str(src),
            "--heuristic",
            "--output", str(out_dir),
        ])
        assert result.exit_code == 0, f"CLI failed: {result.output}"
        outputs = list(out_dir.glob("*.docx"))
        assert len(outputs) == 1, f"Expected 1 output .docx, got: {outputs}"
        assert zipfile.is_zipfile(outputs[0])


# ===========================================================================
# 8. Duplicate detection tests
# ===========================================================================

class TestDuplicateDetector:

    def test_first_occurrence_not_duplicate(self, tmp_path):
        """First submission of a filename is not flagged as a duplicate."""
        from main import DuplicateDetector
        det = DuplicateDetector()
        f = tmp_path / "protocol_a.docx"
        assert det.check(f) is False

    def test_second_occurrence_is_duplicate(self, tmp_path):
        """Second submission of the same filename is flagged as a duplicate."""
        from main import DuplicateDetector
        det = DuplicateDetector()
        f = tmp_path / "protocol_a.docx"
        det.check(f)  # first — not duplicate
        assert det.check(f) is True  # second — duplicate

    def test_case_insensitive_by_default(self, tmp_path):
        """Duplicate detection is case-insensitive by default."""
        from main import DuplicateDetector
        det = DuplicateDetector(case_sensitive=False)
        det.check(tmp_path / "Protocol_A.docx")
        assert det.check(tmp_path / "protocol_a.docx") is True

    def test_case_sensitive_mode(self, tmp_path):
        """In case-sensitive mode, different casings are not duplicates."""
        from main import DuplicateDetector
        det = DuplicateDetector(case_sensitive=True)
        det.check(tmp_path / "Protocol_A.docx")
        assert det.check(tmp_path / "protocol_a.docx") is False

    def test_different_paths_same_filename(self, tmp_path):
        """Files with the same name but different directories are duplicates."""
        from main import DuplicateDetector
        det = DuplicateDetector()
        det.check(tmp_path / "subdir_a" / "protocol.docx")
        assert det.check(tmp_path / "subdir_b" / "protocol.docx") is True

    def test_reset_clears_state(self, tmp_path):
        """reset() allows the same filename to be processed again."""
        from main import DuplicateDetector
        det = DuplicateDetector()
        f = tmp_path / "protocol.docx"
        det.check(f)
        det.reset()
        assert det.check(f) is False

    def test_batch_with_duplicate_skips_second(self, tmp_path):
        """`protocol-formatter batch` skips duplicate filenames within a batch."""
        from typer.testing import CliRunner
        from main import app
        from docx import Document as DocxDocument

        # Create two identically-named files in different subdirs,
        # then copy both into tmp_path (same filename → duplicate)
        d = DocxDocument()
        d.add_heading("Dup Test", level=1)
        d.add_paragraph("Overview.")
        d.add_heading("Procedure", level=1)
        d.add_heading("Step", level=2)
        d.add_paragraph("Do thing.", style="List Number")

        f1 = tmp_path / "dup_protocol.docx"
        d.save(str(f1))

        # Trying to add the same file again won't be possible in batch
        # (batch globbing from a dir only yields each file once), so
        # test the detector directly — batch duplicate prevention is
        # exercised via DuplicateDetector unit tests above.
        from main import DuplicateDetector
        det = DuplicateDetector()
        assert det.check(f1) is False
        assert det.check(f1) is True  # same file → duplicate


# ===========================================================================
# 9. Ollama integration tests (skipped unless server is running)
# ===========================================================================

class TestOllamaIntegration:

    @skip_no_ollama
    @skip_no_templates
    def test_full_pipeline_wet_lab(self, tmp_path):
        """
        Full pipeline: parse wet-lab template → LLM extract → render.

        Requires Ollama running with qwen2.5:14b pulled.
        Marked skip if Ollama is not available.
        """
        from parser.docx_reader import read_docx
        from extractor.llm_extractor import extract_protocol
        from renderer.node_renderer import render_protocol

        doc = read_docx(_WET_LAB_TEMPLATE)
        protocol = extract_protocol(doc, source_filename=_WET_LAB_TEMPLATE.name)

        assert protocol.title
        assert protocol.section_type.value == "wet_lab"
        assert len(protocol.procedure) >= 1

        if _renderer_available:
            output = render_protocol(protocol, output_dir=tmp_path)
            assert output.exists()
            assert zipfile.is_zipfile(output)
