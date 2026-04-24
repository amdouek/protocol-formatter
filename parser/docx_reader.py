"""
parser/docx_reader.py -- python-docx extraction for ProtocolFormatter.

Reads a .docx file and produces a structured ``ParsedDocument`` dataclass that
captures all semantically meaningful content from the source file while
discarding raw presentation details (fonts, colours, sizes).

Preserved content
-----------------
- Paragraph text with bold/italic run properties → **bold**/_italic_ markers
- Heading levels (from paragraph styles)
- Numbered and bulleted lists with nesting depth
- Tables (all cell content, row-by-row)
- Footnotes (collected separately, keyed by their docx footnote id)
- Document core properties (title, author, created/modified dates)

Not preserved
-------------
- Font names, sizes, colours
- Page layout (margins, headers, footers)
- Images and embedded objects
- Comments and tracked changes (changes are ignored; final text is used)
- Styles other than heading level and run bold/italic

Output
------
The ``ParsedDocument`` is consumed by ``extractor/llm_extractor.py``, which
passes the ``full_text`` string and the structured ``paragraphs`` list to the
LLM. The ``footnotes`` dict feeds into ``parser/utils.renumber_footnotes()``.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document as _open_docx          # factory function (used for opening files)
from docx.document import Document as DocxDocument  # actual class (used for type annotations)
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from loguru import logger

from .utils import (
    infer_heading_level,
    normalise_whitespace,
    normalise_cross_references,
    normalise_centrifuge_units,
    runs_to_marked_text,
    extract_author,
    normalise_date,
    infer_section_type,
)


# ---------------------------------------------------------------------------
# Output data structures
# ---------------------------------------------------------------------------

@dataclass
class ParsedRun:
    """A single text run with its inline formatting properties."""
    text: str
    bold: bool = False
    italic: bool = False


@dataclass
class ParsedParagraph:
    """
    A single paragraph extracted from the source document.

    Attributes
    ----------
    text : str
        Full paragraph text with **bold** and _italic_ inline markers applied.
    raw_text : str
        Plain text with no markers (useful for pattern matching).
    style_name : str
        python-docx paragraph style name (e.g. "Heading 1", "Normal", "List Paragraph").
    heading_level : int
        0 = not a heading, 1–3 = H1/H2/H3 as inferred by utils.infer_heading_level.
    list_level : int
        0 = not a list item, 1 = top-level list, 2 = sub-list, etc.
    is_list_item : bool
        True if the paragraph belongs to a numbered or bulleted list.
    is_numbered : bool
        True if the list is a numbered (ordered) list; False for bullets.
    runs : list[ParsedRun]
        The individual runs that make up this paragraph, before merging.
    """
    text: str
    raw_text: str
    style_name: str = ""
    heading_level: int = 0
    list_level: int = 0
    is_list_item: bool = False
    is_numbered: bool = False
    runs: list[ParsedRun] = field(default_factory=list)


@dataclass
class ParsedTable:
    """
    A table extracted from the source document.

    Attributes
    ----------
    rows : list[list[str]]
        Each inner list represents one row; each string is one cell's text,
        whitespace-normalised.
    position : int
        Zero-based index of this table among all block-level elements in the
        document body. Used to reconstruct the original reading order.
    """
    rows: list[list[str]]
    position: int = 0


@dataclass
class ParsedDocument:
    """
    The complete structured output of the docx parser.

    Attributes
    ----------
    paragraphs : list[ParsedParagraph]
        All non-empty paragraphs in document order, including those inside
        table cells (extracted separately in ``tables``).
    tables : list[ParsedTable]
        All tables in document order.
    footnotes : dict[int, str]
        Maps footnote id (int) to footnote text. IDs match those used in the
        document body XML. See utils.renumber_footnotes() to convert to a
        sequential list.
    title : str
        Best-effort protocol title (from document properties or first H1).
    author : str
        Author string (from document properties or byline heuristics).
    date : str | None
        Last-modified or created date in DD/MM/YYYY format, or None.
    section_type : str
        "wet_lab" or "computational" as inferred by utils.infer_section_type.
    full_text : str
        Complete document text as a single newline-separated string, with
        inline markers applied. Passed to the LLM extractor as context.
    source_path : Path
        Absolute path of the source .docx file.
    """
    paragraphs: list[ParsedParagraph] = field(default_factory=list)
    tables: list[ParsedTable] = field(default_factory=list)
    footnotes: dict[int, str] = field(default_factory=dict)
    title: str = ""
    author: str = "ARMI"
    date: Optional[str] = None
    section_type: str = "wet_lab"
    full_text: str = ""
    source_path: Path = field(default_factory=Path)


# ---------------------------------------------------------------------------
# Run extraction
# ---------------------------------------------------------------------------

def _extract_runs(para: DocxParagraph) -> list[ParsedRun]:
    """
    Extract text runs from a python-docx Paragraph, resolving bold and italic
    from both run-level and paragraph-level run properties (rPr/pPr).

    python-docx resolves style inheritance automatically via run.bold and
    run.italic; these can be None (inherit), True, or False.

    Parameters
    ----------
    para : DocxParagraph

    Returns
    -------
    list[ParsedRun]
    """
    parsed_runs: list[ParsedRun] = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        # run.bold / run.italic: True, False, or None (inherit from style)
        bold = bool(run.bold)
        italic = bool(run.italic)
        parsed_runs.append(ParsedRun(text=text, bold=bold, italic=italic))
    return parsed_runs


# ---------------------------------------------------------------------------
# List detection
# ---------------------------------------------------------------------------

def _get_list_info(para: DocxParagraph) -> tuple[bool, int, bool]:
    """
    Determine whether a paragraph is a list item and extract its nesting level.

    Returns
    -------
    tuple[bool, int, bool]
        (is_list_item, list_level, is_numbered)
        list_level is 1-based (0 means not a list item).
    """
    # python-docx exposes paragraph._p (lxml element)
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return False, 0, False

    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        # Check style name as fallback
        style_name = ((para.style.name or "") if para.style else "").lower() # type: ignore[union-attr]
        if "list" in style_name:
            return True, 1, "number" in style_name
        return False, 0, False

    # Extract ilvl (indentation level, 0-based) and numId
    ilvl_el = numPr.find(qn("w:ilvl"))
    ilvl = int(ilvl_el.get(qn("w:val"), 0)) if ilvl_el is not None else 0

    numId_el = numPr.find(qn("w:numId"))
    num_id = int(numId_el.get(qn("w:val"), 0)) if numId_el is not None else 0

    if num_id == 0:
        return False, 0, False

    # Determine ordered vs unordered by inspecting the abstractNumId definition.
    # As a heuristic: paragraph styles containing "number", or ilvl format
    # inspection is expensive; instead check if the style name contains a hint.
    style_name = ((para.style.name or "") if para.style else "").lower() # type: ignore[union-attr]
    is_numbered = (
        "list number" in style_name
        or "numbered" in style_name
        or "ordered" in style_name
    )

    # If style gives no hint, fall back to checking numFmt via document numbering XML
    if not is_numbered:
        try:
            is_numbered = _is_numbered_list(para, num_id, ilvl)
        except Exception:
            # Numbering XML lookup failed; default to numbered (most common in protocols)
            is_numbered = True

    return True, ilvl + 1, is_numbered


def _is_numbered_list(para: DocxParagraph, num_id: int, ilvl: int) -> bool:
    """
    Look up the numFmt for a given numId + ilvl in the document's numbering
    definitions to determine if it's an ordered list.

    Returns True for ordered (decimal, lowerLetter, etc.), False for bullet.
    """
    try:
        doc = para._p.getroottree().getroot()
        # Navigate: w:document → w:body is not the numbering part.
        # The numbering part is a separate XML file; python-docx loads it via
        # para.part.numbering_part
        numbering_part = para.part.numbering_part # type: ignore[attr-defined]
        if numbering_part is None:
            return True

        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = numbering_part._element

        # Find abstractNumId for this numId
        num_el = root.find(
            f".//w:num[@w:numId='{num_id}']",
            namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )
        if num_el is None:
            return True

        abstract_ref = num_el.find(
            "w:abstractNumId",
            namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )
        if abstract_ref is None:
            return True

        abstract_id = abstract_ref.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        )

        abstract_el = root.find(
            f".//w:abstractNum[@w:abstractNumId='{abstract_id}']",
            namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )
        if abstract_el is None:
            return True

        lvl_el = abstract_el.find(
            f".//w:lvl[@w:ilvl='{ilvl}']",
            namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )
        if lvl_el is None:
            return True

        num_fmt_el = lvl_el.find(
            "w:numFmt",
            namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        )
        if num_fmt_el is None:
            return True

        fmt_val = num_fmt_el.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", ""
        )
        return fmt_val != "bullet"

    except Exception as exc:
        logger.debug("_is_numbered_list lookup failed (non-fatal): {}", exc)
        return True


# ---------------------------------------------------------------------------
# Table extraction
# ---------------------------------------------------------------------------

def _extract_table(table: DocxTable, position: int) -> ParsedTable:
    """
    Extract all cell text from a python-docx Table.

    Merged cells (vertical or horizontal spans) are read once at their
    first occurrence; subsequent span cells return empty strings from
    python-docx and are omitted.

    Parameters
    ----------
    table : DocxTable
    position : int

    Returns
    -------
    ParsedTable
    """
    rows: list[list[str]] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            cell_text = " ".join(
                normalise_whitespace(p.text) for p in cell.paragraphs if p.text.strip()
            )
            cells.append(cell_text)
        # Skip entirely blank rows
        if any(c.strip() for c in cells):
            rows.append(cells)
    return ParsedTable(rows=rows, position=position)


# ---------------------------------------------------------------------------
# Footnote extraction
# ---------------------------------------------------------------------------

def _extract_footnotes(doc: DocxDocument) -> dict[int, str]:
    """
    Extract all footnotes from the document's footnotes part.

    python-docx does not expose footnotes through a public API; this function
    accesses the underlying XML directly.

    Returns
    -------
    dict[int, str]
        Maps footnote id (int) to plain-text footnote content.
        IDs 0 and -1 are reserved by Word for separator footnotes and are
        excluded.
    """
    footnotes: dict[int, str] = {}

    try:
        footnotes_part = doc.part.footnotes_part # type: ignore[attr-defined]
    except AttributeError:
        # No footnotes part in this document
        return footnotes

    if footnotes_part is None:
        return footnotes

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for fn_el in footnotes_part._element.findall(".//w:footnote", ns):
        fn_id_str = fn_el.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id"
        )
        try:
            fn_id = int(fn_id_str)
        except (TypeError, ValueError):
            continue

        # Skip separator (id 0) and continuation (id -1) footnotes
        if fn_id <= 0:
            continue

        # Collect all paragraph text in this footnote
        parts: list[str] = []
        for p_el in fn_el.findall(".//w:p", ns):
            for r_el in p_el.findall(".//w:r", ns):
                t_el = r_el.find("w:t", ns)
                if t_el is not None and t_el.text:
                    parts.append(t_el.text)
        text = normalise_whitespace(" ".join(parts))
        if text:
            footnotes[fn_id] = text

    logger.debug("Extracted {} footnote(s)", len(footnotes))
    return footnotes


# ---------------------------------------------------------------------------
# Document properties
# ---------------------------------------------------------------------------

def _extract_core_properties(doc: DocxDocument) -> tuple[str, str, Optional[str]]:
    """
    Extract title, author, and date from the document's core properties.

    Returns
    -------
    tuple[str, str, str | None]
        (title, author, date_string)
    """
    props = doc.core_properties

    raw_title = getattr(props, "title", "") or ""
    raw_author = getattr(props, "author", "") or ""
    raw_modified = getattr(props, "modified", None)
    raw_created = getattr(props, "created", None)

    title = normalise_whitespace(raw_title)
    author = extract_author(normalise_whitespace(raw_author))

    date_str: Optional[str] = None
    for dt_obj in (raw_modified, raw_created):
        if dt_obj is not None:
            try:
                date_str = dt_obj.strftime("%d/%m/%Y")
                break
            except AttributeError:
                pass

    return title, author, date_str


# ---------------------------------------------------------------------------
# Block-level iteration (preserving reading order with tables)
# ---------------------------------------------------------------------------

def _iter_block_items(doc: DocxDocument):
    """
    Yield block-level items (paragraphs and tables) from the document body
    in document order, along with their position index.

    python-docx's doc.paragraphs and doc.tables properties return items in
    document order but separately; iterating the body element directly
    preserves interleaving.

    Yields
    ------
    tuple[int, object]
        (position, item) where item is a DocxParagraph or DocxTable.
    """
    position = 0
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            yield position, DocxParagraph(child, doc)
        elif tag == "tbl":
            yield position, DocxTable(child, doc)
        position += 1


# ---------------------------------------------------------------------------
# Main reader
# ---------------------------------------------------------------------------

def read_docx(source_path: Path) -> ParsedDocument:
    """
    Parse a .docx file and return a ``ParsedDocument``.

    Parameters
    ----------
    source_path : Path
        Absolute or relative path to the .docx file.

    Returns
    -------
    ParsedDocument

    Raises
    ------
    FileNotFoundError
        If ``source_path`` does not exist.
    ValueError
        If the file is not a valid .docx (e.g. an unencrypted .doc).
    """
    source_path = Path(source_path).resolve()

    if not source_path.exists():
        raise FileNotFoundError(f"Source file not found: {source_path}")

    logger.info("Parsing .docx: {}", source_path.name)

    try:
        doc = _open_docx(str(source_path))
    except Exception as exc:
        raise ValueError(
            f"Failed to open '{source_path.name}' as a .docx file. "
            f"If this is a legacy .doc file, use doc_reader.py instead.\n"
            f"Original error: {exc}"
        ) from exc

    # Core properties
    title, author, date_str = _extract_core_properties(doc)
    logger.debug("Core props — title={!r}, author={!r}, date={}", title, author, date_str)

    # Footnotes
    footnotes = _extract_footnotes(doc)

    # Block-level content (paragraphs + tables in order)
    parsed_paragraphs: list[ParsedParagraph] = []
    parsed_tables: list[ParsedTable] = []
    text_lines: list[str] = []

    for position, item in _iter_block_items(doc):
        if isinstance(item, DocxParagraph):
            # Skip empty paragraphs
            if not item.text.strip():
                continue

            runs = _extract_runs(item)
            marked_text = runs_to_marked_text(
                [{"text": r.text, "bold": r.bold, "italic": r.italic} for r in runs]
            )
            raw_text = normalise_whitespace(item.text)
            marked_text = normalise_cross_references(marked_text)
            marked_text = normalise_centrifuge_units(marked_text)

            style_name: str = (item.style.name or "") if item.style else ""
            heading_level = infer_heading_level(raw_text, style_name)
            is_list_item, list_level, is_numbered = _get_list_info(item)

            parsed_para = ParsedParagraph(
                text=marked_text,
                raw_text=raw_text,
                style_name=style_name,
                heading_level=heading_level,
                list_level=list_level,
                is_list_item=is_list_item,
                is_numbered=is_numbered,
                runs=runs,
            )
            parsed_paragraphs.append(parsed_para)

            # Build full_text line
            prefix = ""
            if heading_level == 1:
                prefix = "# "
            elif heading_level == 2:
                prefix = "## "
            elif heading_level == 3:
                prefix = "### "
            elif is_list_item:
                indent = "  " * (list_level - 1)
                marker = f"{list_level}." if is_numbered else "-"
                prefix = f"{indent}{marker} "

            text_lines.append(f"{prefix}{marked_text}")

        elif isinstance(item, DocxTable):
            parsed_table = _extract_table(item, position)
            parsed_tables.append(parsed_table)

            # Represent table in full_text as pipe-separated rows
            for row in parsed_table.rows:
                text_lines.append("| " + " | ".join(row) + " |")

    # Derive title from first H1 if not in core properties
    if not title:
        for para in parsed_paragraphs:
            if para.heading_level == 1 or (
                para.style_name and "title" in para.style_name.lower()
            ):
                from .utils import clean_title
                title = clean_title(para.raw_text)
                break

    # Infer section type
    full_text = "\n".join(text_lines)
    section_type = infer_section_type(full_text)

    result = ParsedDocument(
        paragraphs=parsed_paragraphs,
        tables=parsed_tables,
        footnotes=footnotes,
        title=title,
        author=author,
        date=date_str,
        section_type=section_type,
        full_text=full_text,
        source_path=source_path,
    )

    logger.info(
        "Parsed '{}': {} paragraphs, {} tables, {} footnotes, section_type={}",
        source_path.name,
        len(parsed_paragraphs),
        len(parsed_tables),
        len(footnotes),
        section_type,
    )

    return result